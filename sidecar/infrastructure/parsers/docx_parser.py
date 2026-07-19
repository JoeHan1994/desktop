"""DOCX 文档解析器：提取段落文本并切分为带元数据的文本块。

依赖 python-docx（需在 requirements.txt 中声明）。
"""

from __future__ import annotations

from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter

from ...domain import Chunk, Document
from .markdown_parser import extract_tags

_SPLITTER = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)


def parse_docx_file(
    file_path: Path, docs_root: Path, idf: dict[str, float]
) -> Document:
    """解析单个 DOCX 文件，返回带元数据文本块的 Document。"""
    try:
        import docx  # python-docx 的导入名为 docx
    except ImportError as exc:
        raise ImportError(
            "DOCX 解析需要 python-docx，请运行: pip install python-docx"
        ) from exc

    relative_path = file_path.relative_to(docs_root).as_posix()
    doc = docx.Document(str(file_path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    raw = "\n\n".join(paragraphs)
    lc_chunks = _SPLITTER.create_documents([raw])

    chunks: list[Chunk] = []
    for lc_chunk in lc_chunks:
        text = lc_chunk.page_content
        tags = extract_tags(text, idf)
        chunks.append(
            Chunk(
                text=text,
                source=relative_path,
                metadata={
                    "title": file_path.stem,
                    "file_type": "docx",
                    "has_code": False,
                    "has_steps": False,
                    "has_video": False,
                    "contains_table": False,
                    "tags": ", ".join(tags) if tags else None,
                },
            )
        )
    return Document(source=relative_path, raw=raw, chunks=chunks)
